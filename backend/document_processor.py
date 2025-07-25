import os
from pathlib import Path
import PyPDF2
from docx import Document
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter
from typing import List, Dict, Optional
import magic
import logging
import re
try:
    import spacy
    nlp = spacy.load('en_core_web_sm')
except Exception:
    nlp = None

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=int(os.getenv('CHUNK_SIZE', 1000)),
            chunk_overlap=int(os.getenv('CHUNK_OVERLAP', 200)),
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
        self.allowed_extensions = os.getenv('ALLOWED_EXTENSIONS', '.pdf,.docx,.md,.txt').split(',')
    
    def validate_file(self, file_path: str) -> bool:
        """Validate file type and size"""
        try:
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.allowed_extensions:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Check file size (50MB default limit)
            max_size = 50 * 1024 * 1024  # 50MB in bytes
            if os.path.getsize(file_path) > max_size:
                raise ValueError("File size exceeds maximum limit")
            
            return True
        except Exception as e:
            logger.error(f"File validation error: {str(e)}")
            return False
    
    def process_document(self, file_path: str) -> List[Dict]:
        """Process document and return chunks with metadata, preserving section structure and extracting definitions/tools/materials."""
        try:
            if not self.validate_file(file_path):
                raise ValueError("File validation failed")
            file_ext = Path(file_path).suffix.lower()
            if file_ext == '.pdf':
                text = self._extract_pdf(file_path)
            elif file_ext == '.docx':
                text = self._extract_docx(file_path)
            elif file_ext in ['.md', '.markdown']:
                text = self._extract_markdown(file_path)
            elif file_ext == '.txt':
                text = self._extract_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
            if not text.strip():
                raise ValueError("No text content found in document")
            # Extract sections/headings
            sections = self._extract_sections(text, file_ext)
            # Use sentence-based or semantic splitting if possible
            if nlp:
                doc = nlp(text)
                sentences = [sent.text for sent in doc.sents]
                chunks = self._semantic_split(sentences)
            else:
                chunks = self.text_splitter.split_text(text)
            doc_chunks = []
            for i, chunk in enumerate(chunks):
                section = self._find_section_for_chunk(chunk, sections)
                doc_chunks.append({
                    'text': chunk,
                    'chunk_id': i,
                    'source': Path(file_path).name,
                    'file_type': file_ext,
                    'file_path': file_path,
                    'chunk_size': len(chunk),
                    'section': section,
                    'steps': self._extract_steps(chunk),
                    'safety_notes': self._extract_safety_notes(chunk),
                    'technical_concepts': self._extract_technical_concepts(chunk),
                    'statistical_figures': self._extract_statistical_figures(chunk),
                    'graph_references': self._extract_graph_references(chunk),
                    'definitions': self._extract_definitions(chunk),
                    'tools': self._extract_tools(chunk),
                    'materials': self._extract_materials(chunk),
                    'embedding': None
                })
            logger.info(f"Successfully processed {Path(file_path).name} into {len(chunks)} chunks")
            return doc_chunks
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n" + page_text + "\n"
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")
        return text
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")
        return text
    
    def _extract_markdown(self, file_path: str) -> str:
        """Extract text from Markdown"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return content
        except Exception as e:
            raise ValueError(f"Error reading Markdown: {str(e)}")
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return content
        except Exception as e:
            raise ValueError(f"Error reading text file: {str(e)}")
    
    def _extract_steps(self, text: str) -> List[str]:
        """Extract step-by-step instructions"""
        import re
        
        # Pattern for numbered steps
        numbered_steps = re.findall(r'(?:^|\n)(\d+\.\s+.*?)(?=\n\d+\.|\n\n|\Z)', text, re.MULTILINE | re.DOTALL)
        
        # Pattern for lettered steps
        lettered_steps = re.findall(r'(?:^|\n)([a-z]\.\s+.*?)(?=\n[a-z]\.|\n\n|\Z)', text, re.MULTILINE | re.DOTALL)
        
        # Pattern for bullet points
        bullet_steps = re.findall(r'(?:^|\n)([-*•]\s+.*?)(?=\n[-*•]|\n\n|\Z)', text, re.MULTILINE | re.DOTALL)
        
        all_steps = numbered_steps + lettered_steps + bullet_steps
        return [step.strip() for step in all_steps if step.strip()]
    
    def _extract_safety_notes(self, text: str) -> List[str]:
        """Extract safety-related information"""
        import re
        
        safety_patterns = [
            r'(?i)(?:warning|caution|danger|safety|hazard|risk)[:\s]+(.*?)(?=\n|$)',
            r'(?i)(?:⚠️|🚨|⚡|☢️|☣️)\s*(.*?)(?=\n|$)',
            r'(?i)(?:important|critical|essential)[:\s]+(.*?)(?=\n|$)'
        ]
        
        safety_notes = []
        for pattern in safety_patterns:
            matches = re.findall(pattern, text, re.MULTILINE | re.DOTALL)
            safety_notes.extend([match.strip() for match in matches if match.strip()])
        
        return safety_notes
    
    def _extract_technical_concepts(self, text: str) -> List[str]:
        """Extract technical concepts (e.g., architectures, workflows)"""
        import re
        # Simple keyword-based extraction (expand as needed)
        keywords = [
            r'architecture', r'workflow', r'pipeline', r'algorithm', r'framework',
            r'fusion', r'sensor fusion', r'perception', r'planning', r'control',
            r'neural network', r'CNN', r'LSTM', r'GAN', r'Bayesian', r'Kalman',
            r'object detection', r'tracking', r'localization', r'mapping', r'SLAM'
        ]
        found = set()
        for kw in keywords:
            matches = re.findall(r'\b' + kw + r'\b', text, re.IGNORECASE)
            found.update([m for m in matches])
        return list(found)

    def _extract_statistical_figures(self, text: str) -> List[str]:
        """Extract statistics and figures (e.g., rates, percentages, counts)"""
        import re
        patterns = [
            r'\b\d+\.\d+%\b',         # 12.5%
            r'\b\d+%\b',                # 12%
            r'\b\d+,\d+\b',            # 1,234
            r'\b\d+\.\d+\b',          # 12.34
            r'\b\d+\b',                 # 123
            r'\b(one|two|three|four|five|six|seven|eight|nine|ten)\b (percent|percentage|cases|collisions|accidents|events|failures|successes|instances|samples|trials)'
        ]
        figures = []
        for pat in patterns:
            figures.extend(re.findall(pat, text, re.IGNORECASE))
        # Remove duplicates and flatten tuples
        flat = set()
        for f in figures:
            if isinstance(f, tuple):
                flat.add(' '.join(f))
            else:
                flat.add(f)
        return list(flat)

    def _extract_graph_references(self, text: str) -> List[str]:
        """Extract references to figures, charts, or diagrams"""
        import re
        patterns = [
            r'(Figure|Fig\.|Chart|Diagram|Graph)\s*\d+',
            r'(see|refer to) (Figure|Fig\.|Chart|Diagram|Graph)\s*\d+',
            r'(as shown in|illustrated in) (Figure|Fig\.|Chart|Diagram|Graph)\s*\d+'
        ]
        refs = []
        for pat in patterns:
            refs.extend([m[0] if isinstance(m, tuple) else m for m in re.findall(pat, text, re.IGNORECASE)])
        return list(set(refs))
    
    def _extract_sections(self, text: str, file_ext: str) -> list:
        """Extract section/heading info from text."""
        headings = []
        if file_ext in ['.md', '.markdown', '.txt']:
            for match in re.finditer(r'^(#+|\d+\.|[A-Z][A-Za-z0-9\s]+:)\s+(.+)$', text, re.MULTILINE):
                headings.append({'title': match.group(0).strip(), 'start': match.start()})
        # For DOCX/PDF, could add more advanced logic
        return headings
    def _find_section_for_chunk(self, chunk: str, sections: list) -> str:
        """Find the section/heading for a chunk."""
        # Simple heuristic: find the last heading before the chunk
        for sec in reversed(sections):
            if sec['start'] <= 0:
                return sec['title']
        return sections[0]['title'] if sections else ''
    def _semantic_split(self, sentences: list, max_chunk_size: int = 1000) -> list:
        """Group sentences into chunks of roughly max_chunk_size chars."""
        chunks = []
        current = ''
        for sent in sentences:
            if len(current) + len(sent) > max_chunk_size and current:
                chunks.append(current)
                current = ''
            current += sent + ' '
        if current:
            chunks.append(current.strip())
        return chunks
    def _extract_definitions(self, text: str) -> list:
        """Extract definitions like 'X is ...' or 'X refers to ...'"""
        pattern = r'(\b[A-Z][A-Za-z0-9\s\-]+\b)\s+(is|refers to|means|defined as)\s+(.+?)(\.|$)'
        matches = re.findall(pattern, text)
        return [{'term': m[0].strip(), 'definition': m[2].strip()} for m in matches]
    def _extract_tools(self, text: str) -> list:
        """Extract tool mentions using a keyword list (expand as needed)."""
        tool_keywords = ['wrench', 'screwdriver', 'multimeter', 'oscilloscope', 'soldering iron', 'drill', 'pliers', 'hammer', 'tweezers', 'pipette', 'beaker', 'flask', 'centrifuge', 'pipettor', 'caliper', 'thermometer']
        found = set()
        for kw in tool_keywords:
            if re.search(r'\b' + re.escape(kw) + r'\b', text, re.IGNORECASE):
                found.add(kw)
        return list(found)
    def _extract_materials(self, text: str) -> list:
        """Extract material mentions using a keyword list (expand as needed)."""
        material_keywords = ['acetone', 'ethanol', 'solder', 'wire', 'resistor', 'capacitor', 'hydrochloric acid', 'sodium chloride', 'water', 'oil', 'lubricant', 'glue', 'tape', 'filter', 'membrane', 'buffer', 'agarose', 'gel']
        found = set()
        for kw in material_keywords:
            if re.search(r'\b' + re.escape(kw) + r'\b', text, re.IGNORECASE):
                found.add(kw)
        return list(found)
    
    def get_document_summary(self, file_path: str) -> Dict:
        """Get document summary without full processing"""
        try:
            file_stat = os.stat(file_path)
            file_ext = Path(file_path).suffix.lower()
            
            return {
                'filename': Path(file_path).name,
                'file_type': file_ext,
                'file_size_mb': round(file_stat.st_size / (1024 * 1024), 2),
                'last_modified': file_stat.st_mtime,
                'is_supported': file_ext in self.allowed_extensions
            }
        except Exception as e:
            logger.error(f"Error getting document summary: {str(e)}")
            return {}
